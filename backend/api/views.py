from django.shortcuts import render
from rest_framework import generics, status
from .serializers import SenadorSerializer, CreateSenadorSerializer, ComisionSerializer, MocionSerializer, EventSerializer, LeyesEnTramitacionSerializer, EnmiendaSerializer, PreguntaSerializer, PlenoSerializer
from .models import Senador, Comision, Mocion, Event, LeyesEnTramitacion, Enmienda, Pregunta, Pleno
from rest_framework.views import APIView
from rest_framework.response import Response
import requests
from bs4 import BeautifulSoup
import datetime
import docx
import os
from django.core.files.storage import FileSystemStorage
from django.shortcuts import HttpResponse
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.sessions.backends.db import SessionStore
from pathlib import Path
BASE_DIR = Path(__file__).resolve().parent.parent
s = SessionStore()


class Login(APIView):
    def post(self, req, format=None):
        username = req.data.get('username')
        password = req.data.get('password')
        user = authenticate(req, username=username, password=password)
        if user is not None:
            login(req, user)
            return Response(req.session, status=status.HTTP_200_OK)
        else:
            return Response('Usuario o contraseña incorrectos', status=status.HTTP_401_UNAUTHORIZED)


class Logout(APIView):
    def post(self, req, format=None):
        print(req)
        logout(req)
        return Response("Logged out", status=status.HTTP_200_OK)

#################################################
# SENADORES
#################################################


class GetSenadoresView(generics.ListAPIView):
    # queryset = Senador.objects.all()
    serializer_class = SenadorSerializer
    # @login_required

    def get(self, req, format=None):
        senadores = Senador.objects.all()
        return Response(senadores.values(), status=status.HTTP_200_OK)


class GetSenador(APIView):
    serializer_class = SenadorSerializer
    lookup_url_kwarg = 'id'

    def get(self, req, format=None):
        id = req.GET.get(self.lookup_url_kwarg)
        if id != None:
            senador = Senador.objects.filter(id=id)
            if len(senador) > 0:
                data = SenadorSerializer(senador[0]).data
                return Response(data, status=status.HTTP_200_OK)
            return Response({'No se ha encontrado al Senador'}, status=status.HTTP_404_NOT_FOUND)
        return Response({'No se ha encontrado el ID'}, status=status.HTTP_400_BAD_REQUEST)


class CreateSenadorView(APIView):
    serializer_class = CreateSenadorSerializer
    # define PUT, POST, GET methods

    def post(self, req, format=None):
        serializer = self.serializer_class(data=req.data)
        if serializer.is_valid():
            name = serializer.data.get('name')
            last_name = serializer.data.get('last_name')
            senador = Senador(name=name, last_name=last_name)
            senador.save()
            return Response(status=status.HTTP_200_OK)
        return Response(status=status.HTTP_400_BAD_REQUEST)


class DeleteSenadorView(APIView):
    @login_required
    def delete(self, req, id, format=None):
        senador = Senador.objects.get(id=id)
        # print(senador)
        senador.delete()
        return Response(status=status.HTTP_200_OK)

#################################################
# COMISIONES
#################################################


class GetComisionesView(APIView):
    serializer_class = ComisionSerializer

    def get(self, req, format=None):
        comisiones = Comision.objects.all()
        # print(comisiones.values().)
        return Response(comisiones.values('id', 'evento__subject', 'evento__startTime', 'senador__name', 'senador__last_name').order_by('-evento__startTime'), status=status.HTTP_200_OK)


class DeleteComisionView(APIView):
    def delete(self, req, id, format=None):
        evento = Event.objects.filter(comision=id)
        evento.delete()
        return Response(status=status.HTTP_200_OK)


class GetComisionView(APIView):
    serializer_class = SenadorSerializer
    lookup_url_kwarg = 'id'

    def get(self, req, format=None):
        id = req.GET.get(self.lookup_url_kwarg)
        if id != None:
            comision = Comision.objects.filter(id=id)
            return Response(comision.values('evento__subject', 'evento__startTime', 'url', 'senador__name', 'senador__last_name'), status=status.HTTP_200_OK)
        return Response({'No se ha encontrado el ID'}, status=status.HTTP_400_BAD_REQUEST)


#################################################
# MOCIONES
#################################################

class PostMocion(APIView):
    serializer_class = MocionSerializer

    def post(self, req, format=None):
        xml = requests.get(req.data.get('url')).text
        soup_pdf = BeautifulSoup(xml, 'lxml').find(
            attrs={'class': 'icon-registro'}).find('a').get('href')
        mocion_data = requests.get('https://www.senado.es'+soup_pdf).text
        soup_mocion = BeautifulSoup(mocion_data, 'lxml').find(
            attrs={'class': 'icon-pdf'}).get('href')
        #####
        data = {
            'title': req.data.get('title'),
            'url': req.data.get('url'),
            'exp': req.data.get('exp'),
            'pleno': int(req.data.get('pleno_id')),
            'pdf': soup_mocion
        }
        print(data)
        serializer = self.serializer_class(data=data)
        
        if serializer.is_valid():
            serializer.save()
            return Response(status=status.HTTP_200_OK)
        print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class GetMocionesPleno(APIView):

    def get(self, req, id, format=None):
        mociones = Mocion.objects.filter(pleno__id=id)
        print(mociones)
        return Response(mociones.values(), status=status.HTTP_200_OK)

class GetMocion(APIView):
    serializer_class = MocionSerializer
    # lookup_url_kwarg = 'comision_id'

    def get(self, req, id, fromat=None):
        mocion = Mocion.objects.filter(id=id)
        data = self.serializer_class(mocion[0]).data
        return Response(data, status=status.HTTP_200_OK)

class GetMocionDetail(APIView):
    serializer_class = MocionSerializer
    # lookup_url_kwarg = 'comision_id'

    def get(self, req, id, fromat=None):
        mocion = Mocion.objects.filter(id=id)
        data = self.serializer_class(mocion[0]).data
        url = data.get('url')
        xml = requests.get(url).text
        soup = BeautifulSoup(xml, 'lxml').find(
            attrs={'class': 'icon-registro'}).find('a').get('href')
        mocion_data = requests.get('https://www.senado.es'+soup).text
        soup_mocion = BeautifulSoup(mocion_data, 'lxml').find(
            attrs={'class': 'icon-pdf'}).get('href')
        # print(soup_mocion)
        return Response(soup_mocion, status=status.HTTP_200_OK)


class GetMocionesView(APIView):
    serializer_class = MocionSerializer
    lookup_url_kwarg = 'comision_id'

    def get(self, req, format=None):
        id = req.GET.get(self.lookup_url_kwarg)
        mociones = Mocion.objects.filter(comision_id=id)
        return Response(mociones.values('id', 'title', 'url', 'enmienda', 'discurso', 'voto', 'pdf'), status=status.HTTP_200_OK)


class GetMocionesPendientesView(APIView):
    serializer_class = MocionSerializer

    def get(self, req, format=None):
        date = datetime.datetime.today().strftime("%A")
        self.include_mociones()
        mociones = Mocion.objects.all()
        return Response(mociones.values('title', 'url', 'comision_global', 'comision__evento__startTime', 'comision__evento__subject', 'pleno'), status=status.HTTP_200_OK)

    def include_mociones(self):
        constitucional = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011003&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        exteriores = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?legis=14&id=S011001&tab=2').text
        cooperacion = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011019&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        sanidad = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011012&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        cultura = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011007&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        entidades = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011017&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        migraciones = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011013&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        familia = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S012008&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        defensa = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011004&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        interior = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S011010&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text
        incompatibilidades = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/sesionescomision/detallecomisiones/iniciativasentramitacion/index.html?id=S012006&legis=14&esMixta=N&seccionN1=1318&seccionActual=Composicion').text

        self.get_mocion(constitucional, "Constitucional")
        self.get_mocion(exteriores, "Exteriores")
        self.get_mocion(cooperacion, "Cooperacion Internacional")
        self.get_mocion(sanidad, "Sanidad y Consumo")
        self.get_mocion(cultura, "Cultura y Deporte")
        self.get_mocion(entidades, "Entidades Locales")
        self.get_mocion(
            migraciones, "Trabajo,Inclusión, Seguridad Social y Migraciones")
        self.get_mocion(familia, "Derechos de Familia")
        self.get_mocion(defensa, "Defensa")
        self.get_mocion(interior, "Interior")
        self.get_mocion(incompatibilidades, "Incompatibilidades")

    def get_mocion(self, data, com_str):
        soup = BeautifulSoup(data, 'lxml')
        tags = soup.findAll(attrs={'class': 'text_c2'})
        for tag in tags:
            moc_url = "https://www.senado.es"+tag.get('href')
            data = {
                'title': str(tag.contents[0]),
                'url': moc_url,
                'comision_global': com_str,
            }
            serializer = self.serializer_class(data=data)
            if serializer.is_valid():
                serializer.save()


class ChangeVote(APIView):
    def post(self, req, format=None):
        id = req.data.get('id')
        vote = req.data.get('vote')
        queryset = Mocion.objects.filter(id=id)
        if queryset.exists():
            mocion = queryset[0]
            mocion.voto = int(vote)
            mocion.save()
            return Response(status=status.HTTP_200_OK)
        else:
            return Response(status=status.HTTP_400_BAD_REQUEST)


#################################################
# EVENTOS
#################################################
class PostEventView(APIView):
    serializer_class = EventSerializer
    comision_serializer_class = ComisionSerializer
    mocion_serializer_class = MocionSerializer
    pleno_serializer = PlenoSerializer

    def post(self, req, format=None):
        start_date_time_str = req.data.get('startTime')
        start_date_time_obj = datetime.datetime.strptime(
            start_date_time_str, "%Y-%m-%dT%H:%M:%S.%fZ")
        end_date_time_str = req.data.get('endTime')
        end_date_time_obj = datetime.datetime.strptime(
            end_date_time_str, "%Y-%m-%dT%H:%M:%S.%fZ")
        data = {
            'subject': req.data.get('subject'),
            'startTime': start_date_time_obj,
            'endTime': end_date_time_obj,
            'isAllDayEvent': req.data.get('isAllDayEvent'),
            'tipo': req.data.get('type')
        }
        serializer = self.serializer_class(data=data)
        if serializer.is_valid():
            serializer.save()
            if serializer.data.get('tipo') == 'comision':
                self.post_comision(serializer.data.get('id'), req.data.get(
                    'senador_id'), req.data.get('url'), serializer.data.get('subject'))
            elif serializer.data.get('tipo') == 'pleno':
                self.post_pleno(serializer.data.get('id'), req.data.get('url'))
            return Response(status=status.HTTP_200_OK)
        return Response(status=status.HTTP_400_BAD_REQUEST)

    def post_pleno(self, event_id, url, format=None):
        data = {
            'url': url,
            'evento': event_id
        }
        serializer = self.pleno_serializer(data=data)
        if serializer.is_valid():
            print("pleno creado")
            serializer.save()
            return Response(status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def post_comision(self, event_id, senador_id, url, subject, format=None):
        data = {
            'url': url,
            'senador': senador_id,
            'evento': event_id
        }
        serializer = self.comision_serializer_class(data=data)
        if serializer.is_valid():
            serializer.save()
            print("Comision Creada")
            id = serializer.data.get('id')
            url = serializer.data.get('url')
            self.post_mocion(url, id, subject)
            return Response(status=status.HTTP_200_OK)

        return Response(status=status.HTTP_400_BAD_REQUEST)

    def post_mocion(self, url, id, subject):
        datos = requests.get(url).text
        soup = BeautifulSoup(datos, 'html.parser')
        tags = soup.findAll(attrs={'class': 'text_c2'})
        exps = soup.find_all('span', attrs={'class': 'new-line'})
        exps_list = []
        # print(exps)
        i = -1
        for exp in exps:
            exps_list.append(exp.text)

        for tag in tags:
            i += 1
            moc_url = "https://www.senado.es"+tag.get('href')
            ####
            xml = requests.get(moc_url).text
            soup_pdf = BeautifulSoup(xml, 'lxml').find(
                attrs={'class': 'icon-registro'}).find('a').get('href')
            mocion_data = requests.get('https://www.senado.es'+soup_pdf).text
            soup_mocion = BeautifulSoup(mocion_data, 'lxml').find(
                attrs={'class': 'icon-pdf'}).get('href')
            #####
            data = {
                'title': tag.contents[0],
                'url': moc_url,
                'comision': id,
                'exp': exps_list[i],
                'comision_global': subject,
                'pdf': soup_mocion
            }
            serializer = self.mocion_serializer_class(data=data)
            if serializer.is_valid():
                serializer.save()
        return Response(status=status.HTTP_200_OK)


class GetEventsView(APIView):
    serializer_class = EventSerializer

    def get(self, req, format=None):
        events = Event.objects.all()
        return Response(events.values(), status=status.HTTP_200_OK)


class DeleteEventView(APIView):
    serializer_class = EventSerializer

    def delete(self, req, id, format=None):
        event = Event.objects.get(id=id)
        event.delete()
        return Response(status=status.HTTP_200_OK)

#################################################
# LEYES EN TRAMITACIÓN
#################################################


class GetLeyesView(APIView):
    serializer_class = LeyesEnTramitacionSerializer

    def get(self, req, format=None):
        self.include_leyes()
        leyes = LeyesEnTramitacion.objects.all()
        return Response(leyes.values().order_by('-created_at'))

    def include_leyes(self):
        leyes = requests.get(
            'https://www.senado.es/web/actividadparlamentaria/actualidad/leyes/entramitacion/index.html').text
        soup = BeautifulSoup(leyes, 'html.parser')
        tags = soup.findAll(attrs={'class': 'destacado_azul negrita'})
        for tag in tags:
            moc_url = "https://www.senado.es"+tag.get('href')
            data = {
                'name': str(tag.contents[0]),
                'url': moc_url
            }
            # print(data)
            serializer = self.serializer_class(data=data)
            if serializer.is_valid():
                serializer.save()

#################################################
# ENMIENDAS
#################################################


class PostEnmiendaView(APIView):
    serializer_class = EnmiendaSerializer

    def post(self, req, format=None):
        mocion_id = req.data.get('mocion_id')
        # print(mocion_id)
        tipo = req.data.get('tipo')
        enmienda = req.data.get('enmienda')
        senador = req.data.get('senador')
        # exp = req.data.get('exp')
        # name = exp+'_'+tipo
        # print(req.data)
        sec_path = 'media/comisiones/'
        path = os.path.join(BASE_DIR, sec_path)
        dirname = os.path.join(os.path.dirname(path), 'Plantilla_VOX.docx')
        mocion = Mocion.objects.filter(id=mocion_id)
        moc = MocionSerializer(mocion[0]).data.get('title')
        expediente = MocionSerializer(mocion[0]).data.get('exp').strip('()')
        name = expediente+'_'+tipo+'.docx'
        data = {
            'title': name,
            'url': os.path.join(os.path.dirname(path), name),
            'mocion': mocion_id
        }
        serializer = self.serializer_class(data=data)
        # print(serializer.is_valid())
        if serializer.is_valid():
            serializer.save()
            doc = docx.Document(dirname)
            doc.paragraphs[2].runs[1].text = senador
            doc.paragraphs[2].runs[3].text = 'la siguiente enmienda de '+tipo
            doc.paragraphs[2].runs[6].text = expediente
            doc.paragraphs[2].runs[9].text = 'Partido político'
            doc.paragraphs[2].runs[11].text = moc
            doc.paragraphs[6].runs[1].text = datetime.date.today().strftime(
                "%d/%m/%Y")
            doc.paragraphs[4].text = enmienda
            doc.save(os.path.join(os.path.dirname(path), name))
            return Response(status=status.HTTP_200_OK)
        # print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class GetEnmiendaView(APIView):
    serializer_class = EnmiendaSerializer
    lookup_url_kwarg = 'mocion_id'

    def get(self, req, format=None):
        id = req.GET.get(self.lookup_url_kwarg)
        enmienda = Enmienda.objects.filter(mocion_id=id)
        # print(enmienda)
        return Response(enmienda.values('id', 'mocion_id', 'mocion__discurso', 'title', 'url'), status=status.HTTP_200_OK)


class DownloadEnmiendaView(APIView):
    serializer_class = EnmiendaSerializer
    lookup_url_kwarg = 'mocion_id'

    def get(self, req, format=None):
        sec_path = 'media/comisiones/'
        path = os.path.join(BASE_DIR, sec_path)
        mocion_id = req.GET.get(self.lookup_url_kwarg)
        enmienda = Enmienda.objects.filter(mocion=mocion_id)
        if len(enmienda) > 0:
            data = self.serializer_class(enmienda[0]).data
        name = data.get('title')
        file_path = os.path.join(os.path.dirname(path), name)
        # print(file_path)
        with open(file_path, 'rb') as fh:
            response = HttpResponse(
                fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'attachment; filename=' + \
                os.path.basename(file_path)
            return response


class DeleteEnmiendaView(APIView):
    serializer_class = EnmiendaSerializer

    def delete(self, req, id, format=None):
        enmienda = Enmienda.objects.filter(id=id)
        serializer = self.serializer_class(enmienda[0]).data
        fs = FileSystemStorage()
        fs.delete(serializer.get('url'))
        enmienda.delete()
        return Response(status=status.HTTP_200_OK)


class PostDiscursoView(APIView):
    serializer_class = MocionSerializer

    def post(self, req, format=None):
        mocion_id = req.data.get('id')
        doc = req.data['selectedFile']
        # name = req.data['name']
        queryset = Mocion.objects.filter(id=mocion_id)
        if queryset.exists():
            mocion = queryset[0]
            mocion.discurso = doc
            mocion.save(update_fields=['discurso'])
            return Response(status=status.HTTP_202_ACCEPTED)
        else:
            return Response(status=status.HTTP_400_BAD_REQUEST)


class DownloadDiscursoView(APIView):
    serializer_class = MocionSerializer
    lookup_url_kwarg = 'id'

    def get(self, req, format=None):
        id = req.GET.get(self.lookup_url_kwarg)
        mocion = Mocion.objects.filter(id=id)
        if len(mocion) > 0:
            data = self.serializer_class(mocion[0]).data
        print(data)
        sec_path = data['discurso'].split('/')
        file_path = os.path.join(BASE_DIR, '/'.join(sec_path[1:]))
        with open(file_path, 'rb') as fh:
            response = HttpResponse(
                fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'attachment; filename=' + \
                os.path.basename(file_path)
            return response


class DeleteDiscursoView(APIView):
    def delete(self, req, id, format=None):
        pass

#################################################
# SESSION
#################################################


class UserLogged(APIView):
    pass


#################################################
# PREGUNTA
#################################################
class PostPregunta(APIView):
    serializer_class = PreguntaSerializer

    def post(self, req, format=None):
        senador_id = req.data.get('id')
        doc = req.data['selectedFile']
        data = {
            'senador': senador_id,
            'document': doc
        }
        serializer = self.serializer_class(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(status=status.HTTP_202_ACCEPTED)
        print(serializer.errors)
        return Response(status=status.HTTP_400_BAD_REQUEST)


class GetPreguntas(APIView):
    # @login_required
    def get(self, req, format=None):
        preguntas = Pregunta.objects.all()
        return Response(preguntas.values('id', 'created_at', 'senador__name', 'senador__last_name', 'document', 'status'), status=status.HTTP_200_OK)


class DownloadPreguntaView(APIView):
    serializer_class = PreguntaSerializer
    lookup_url_kwarg = 'id'

    def get(self, req, format=None):
        id = req.GET.get(self.lookup_url_kwarg)
        pregunta = Pregunta.objects.filter(id=id)
        if len(pregunta) > 0:
            data = self.serializer_class(pregunta[0]).data
        print(data)
        sec_path = data['document'].split('/')
        file_path = os.path.join(BASE_DIR, '/'.join(sec_path[1:]))
        with open(file_path, 'rb') as fh:
            response = HttpResponse(
                fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'attachment; filename=' + \
                os.path.basename(file_path)
            return response


class DeletePreguntaView(APIView):
    serializer_class = PreguntaSerializer

    def delete(self, req, id, format=None):
        pregunta = Pregunta.objects.filter(id=id)
        serializer = self.serializer_class(pregunta[0]).data
        fs = FileSystemStorage()
        sec_path = serializer['document'].split('/')
        file_path = os.path.join(BASE_DIR, '/'.join(sec_path[1:]))
        fs.delete(file_path)
        pregunta.delete()
        return Response(status=status.HTTP_200_OK)


#################################################
# PLENO
# ###############################################

class GetPlenosView(APIView):
    def get(self, req, format=None):
        plenos = Pleno.objects.all()
        return Response(plenos.values('id', 'evento__startTime', 'evento__subject', 'evento__id', 'url'), status=status.HTTP_200_OK)


class GetPlenoView(APIView):
    lookup_url_kwarg = 'id'

    def get(self, req, id, format=None):
        # id = req.GET.get(self.lookup_url_kwarg)
        if id != None:
            pleno = Pleno.objects.filter(id=id)
            return Response(pleno.values('evento__startTime', 'url'), status=status.HTTP_200_OK)
        return Response({'No se ha encontrado el ID'}, status=status.HTTP_400_BAD_REQUEST)
